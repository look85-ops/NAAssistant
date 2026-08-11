# -*- coding: utf-8 -*-
"""Insert slides into free markers. Usage: change START, END, DOCX, OUT."""
import os, re
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# === CONFIG ===
DOCX = r'C:\Users\marcenuk\Desktop\М1Д1_работа2\Сценарий_М1Д1_26_working.docx'
OUT = r'C:\Users\marcenuk\Desktop\М1Д1_работа2\Сценарий_М1Д1_26_working.docx'
PPTX = r'C:\Users\marcenuk\Desktop\М1Д1_работа2\Управляющая_презентация_М1Д1_26FV.pptx'
TDIR = r'C:\Users\marcenuk\Desktop\slide_exports'
START, END = 189, 254
W, H = 8.8, 5.5

ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
bl = 'http://schemas.openxmlformats.org/drawingml/2006/main'
MRK = re.compile(r'^Слайд(\s*№\s*\d+)?\s*$')

# === Export missing slides ===
ex = set(os.listdir(TDIR))
miss = [i for i in range(START, END + 1) if 'slide_%03d.png' % i not in ex]
if miss:
    print('Exporting %d slides...' % len(miss))
    import win32com.client
    pp = win32com.client.Dispatch('PowerPoint.Application')
    pp.Visible = True
    prs = pp.Presentations.Open(PPTX, WithWindow=False)
    for sn in miss:
        prs.Slides(sn).Export(os.path.join(TDIR, 'slide_%03d.png' % sn), 'PNG')
    prs.Close()
    pp.Quit()

SF = {}
for i in range(START, END + 1):
    p = os.path.join(TDIR, 'slide_%03d.png' % i)
    if os.path.exists(p):
        SF[i] = p

# === Find free markers ===
doc = Document(DOCX)
free = []
for ti, t in enumerate(doc.tables):
    for ri, row in enumerate(t.rows):
        if len(row.cells) <= 4:
            continue
        cell = row.cells[4]
        tc = cell._tc
        ps = list(tc.findall('.//{%s}p' % ns['w']))
        for pi, p_elem in enumerate(ps):
            txt = ''.join(t.text or '' for t in p_elem.findall('.//{%s}t' % ns['w'])).strip()
            if not MRK.match(txt):
                continue
            has_img = (pi + 1 < len(ps)) and (
                ps[pi + 1].findall('.//{%s}inline' % wp) or
                ps[pi + 1].findall('.//{%s}blip' % bl))
            if not has_img:
                free.append((ti, ri, p_elem, txt[:40]))

print('Free markers: %d, taking %d (slides %d-%d)' % (len(free), END - START + 1, START, END))

# === Insert ===
from collections import defaultdict
bc = defaultdict(list)
for sn, (ti, ri, pe, tx) in enumerate(free[:END - START + 1], start=START):
    bc[(ti, ri)].append((pe, sn, tx))

ins = 0
for ti, ri in sorted(bc.keys(), reverse=True):
    items = bc[(ti, ri)]
    cell = doc.tables[ti].rows[ri].cells[4]
    tc = cell._tc
    cps = list(tc.findall('.//{%s}p' % ns['w']))
    
    posd = []
    for pe, sn, tx in items:
        pos = None
        for pi, cp in enumerate(cps):
            if cp is pe:
                pos = pi
                break
        if pos is None:
            for pi, cp in enumerate(cps):
                ct = ''.join(t.text or '' for t in cp.findall('.//{%s}t' % ns['w'])).strip()[:40]
                if ct == tx:
                    hi = (pi + 1 < len(cps)) and (
                        cps[pi + 1].findall('.//{%s}inline' % wp) or
                        cps[pi + 1].findall('.//{%s}blip' % bl))
                    if not hi:
                        pos = pi
                        pe = cp
                        break
        if pos is not None:
            posd.append((pos, pe, sn, tx))
    
    posd.sort(key=lambda x: x[0], reverse=True)
    
    for pos, pe, sn, tx in posd:
        if sn not in SF:
            continue
        np = cell.add_paragraph()
        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = np.add_run()
        r.add_picture(SF[sn], width=Cm(W), height=Cm(H))
        tc.remove(np._element)
        pe.addnext(np._element)
        ins += 1

# Show summary per row
from collections import Counter
row_counts = Counter()
for ti, ri, pe, tx in free[:END - START + 1]:
    row_counts[(ti, ri)] += 1
print('Inserted: %d' % ins)
for (ti, ri), cnt in sorted(row_counts.items()):
    print('  T%d R%d: %d slides' % (ti, ri, cnt))
print('Remaining free: %d' % (len(free) - ins))

doc.save(OUT)
print('Saved: %s' % OUT)
