# Конвертирует .md → .docx для Сигнала.
# Запуск: C:\Users\marcenuk\AppData\Local\Python\pythoncore-3.14-64\python.exe scripts/convert_md_to_docx.py <input.md> [output.docx]
# По умолчанию: knowledge/signal/YYYY-MM-DD.md → knowledge/signal/YYYY-MM-DD.docx

import sys, re, os
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    os.system("pip install python-docx")
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('- **'):
            rest = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            bold_end = rest.find('**', 2)
            if bold_end != -1:
                p.add_run(rest[2:bold_end]).bold = True
                after = rest[bold_end+2:]
                if after.startswith('.'):
                    after = after[1:]
                p.add_run(after)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('─' * 40)
        elif line.startswith('_'):
            p = doc.add_paragraph()
            p.add_run(line).italic = True
    
    doc.save(docx_path)
    print(f"OK: {docx_path}")

if __name__ == '__main__':
    if len(sys.argv) >= 2:
        src = sys.argv[1]
        dst = sys.argv[2] if len(sys.argv) >= 3 else src.replace('.md', '.docx')
    else:
        # Default: latest signal
        signal_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'knowledge', 'signal')
        src = os.path.join(signal_dir, '2026-07-22.md')
        dst = src.replace('.md', '.docx')
    
    convert(src, dst)
