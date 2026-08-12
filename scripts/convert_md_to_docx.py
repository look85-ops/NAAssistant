# Конвертирует .md → .docx (Сигнал + уроки L&DxAI).
# Запуск: C:\Users\marcenuk\AppData\Local\Python\pythoncore-3.14-64\python.exe scripts/convert_md_to_docx.py <input.md> [output.docx]

import sys, re
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_runs(paragraph, text):
    """Разбирает **жирный** и `код` фрагменты, добавляет в параграф."""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif part:
            paragraph.add_run(part)

def parse_cells(line):
    """Вытаскивает ячейки из markdown-строки таблицы."""
    return [c.strip() for c in line.strip('|').split('|')]

def is_table_separator(line):
    """Строка-разделитель markdown-таблицы: | --- | --- |"""
    return re.match(r'^\|[\s\-:|]+\|$', line)

def add_table(doc, header_row, body_rows):
    """Создаёт таблицу docx с шапкой и строками данных."""
    ncols = len(header_row)
    nrows = 1 + len(body_rows)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'
    for j, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        add_runs(p, cell_text)
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
    for i, row in enumerate(body_rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            add_runs(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(9)
    doc.add_paragraph()  # отступ после таблицы

def is_table_row(line):
    """Является ли строка строкой markdown-таблицы (не разделитель)."""
    return line.startswith('|') and line.endswith('|') and not is_table_separator(line)

def convert(md_path, docx_path):
    doc = Document()

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Двухпроходный сбор: сначала группируем таблицы
    table_buffer = []      # временно накапливаем строки таблицы
    in_code_block = False

    for raw in lines:
        line = raw.rstrip()

        # Пропуск кодовых блоков — внутрь не лезем
        if line.startswith('```'):
            in_code_block = not in_code_block
            # Добавляем разделитель
            p = doc.add_paragraph()
            p.add_run('```').font.size = Pt(8)
            continue
        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            continue

        if not line.strip():
            # Пустая строка: сбрасываем буфер таблицы, если он есть
            if table_buffer:
                header = parse_cells(table_buffer[0])
                body = [parse_cells(r) for r in table_buffer[1:]]
                add_table(doc, header, body)
                table_buffer = []
            continue

        # Таблица: накапливаем строки
        if is_table_row(line):
            table_buffer.append(line)
            continue
        elif is_table_separator(line):
            # Разделитель — его место между header и body уже учтено:
            # если буфер пуст или содержит только header — ок, иначе сбрасываем буфер
            if table_buffer:
                # разделитель между header и body — просто пропускаем
                pass
            continue

        # Не таблица — если был буфер, сбрасываем
        if table_buffer:
            header = parse_cells(table_buffer[0])
            body = [parse_cells(r) for r in table_buffer[1:]]
            add_table(doc, header, body)
            table_buffer = []

        # Обычная обработка строки
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('---') or line.startswith('***'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('─' * 40)
        elif line.startswith('*') and line.endswith('*') and len(line) > 2:
            p = doc.add_paragraph()
            p.add_run(line[1:-1]).italic = True
        elif line.startswith('_') and line.endswith('_') and len(line) > 2:
            p = doc.add_paragraph()
            p.add_run(line[1:-1]).italic = True
        elif re.match(r'^\s*[-*] ', line):
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, line.lstrip()[2:])
        elif re.match(r'^\s*\d+\.\s+', line):
            p = doc.add_paragraph(style='List Number')
            add_runs(p, re.sub(r'^\s*\d+\.\s+', '', line))
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
        else:
            p = doc.add_paragraph()
            add_runs(p, line)

    # Финал: сбросить оставшийся буфер
    if table_buffer:
        header = parse_cells(table_buffer[0])
        body = [parse_cells(r) for r in table_buffer[1:]]
        add_table(doc, header, body)

    doc.save(docx_path)
    print(f"OK: {docx_path}")

if __name__ == '__main__':
    if len(sys.argv) >= 2:
        src = sys.argv[1]
        dst = sys.argv[2] if len(sys.argv) >= 3 else src.replace('.md', '.docx')
        convert(src, dst)
    else:
        print("Usage: python convert_md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)