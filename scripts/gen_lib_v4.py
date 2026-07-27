import os
import re
from docx import Document
from docx.shared import Pt


def load_deep_protocols(src_path: str):
    with open(src_path, encoding='utf-8') as f:
        html = f.read()
    m = re.search(r'const DEEP_PROTOCOLS\s*=\s*\[(.*?)\];', html, re.DOTALL)
    if not m:
        return []
    body = m.group(1)
    protocols = []
    idx = 0

    def extract(field: str, text: str) -> str:
        m2 = re.search(rf"{field}:'((?:[^'\\]|\\.|'')*)'", text)
        if m2:
            val = m2.group(1).replace("''", "'")
            return val.replace('\\n', '\n')
        return ''

    while True:
        start = body.find("{ id:'", idx)
        if start == -1:
            break
        end = body.find("},\n    {", start)
        if end == -1:
            end = body.find("}\n  ]", start)
        if end == -1:
            end = len(body)
        block = body[start:end]
        prot = {k: extract(k, block) for k in ('id', 'cat', 'title', 'desc', 'prompt')}
        if prot.get('id'):
            protocols.append(prot)
        idx = end + 1
    return protocols


def build_docx(protocols, out_path: str):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Intro: positioning + temperature guidance
    doc.add_heading('Prompt Library v5 — ядро протоколов для методиста', 0)
    intro1 = (
        'Этот набор — не «всё обо всём», а ядро протоколов, которые дают несоразмерную пользу методисту. '
        'Каждый из них закрывает реальную боль: быстро вытащить знание из эксперта, превратить сопротивление в практику, '
        'связать кейсы и тесты с целями, или разложить размытые компетенции в наблюдаемое поведение.'
    )
    intro2 = (
        'Мы оставили только то, что трудно найти в открытых библиотеках и что работает в реальных корпоративных проектах. '
        'Минимум теории, максимум действий и проверок: у каждого протокола есть «когда не использовать» и валидация, '
        'чтобы не получить красивый, но бесполезный текст.'
    )
    intro3 = (
        'Температура — насколько модель позволяет себе креативность и вариативность. '
        '0.1–0.2: строго и предсказуемо (цели, чек-листы, рубрики, тесты). '
        '0.3–0.4: умеренная вариативность (аналитика, проектирование). '
        '0.5–0.7: креатив и поиск вариантов (кейсы, сценарии, сторителлинг). '
        'Выше 0.7 — только осознанно: брейншторм и неожиданные ходы; риск мусора растёт.'
    )
    doc.add_paragraph(intro1)
    doc.add_paragraph(intro2)
    doc.add_paragraph(intro3)
    recs = {
        'Диагностика': 'Температура 0.2–0.3: минимум «творчества», максимум точности и перепроверки.',
        'Проектирование': 'Температура 0.3–0.4: идеи допустимы, но держим рамку и валидируем.',
        'Коммуникации': 'Температура 0.3–0.5: письма и повестки — 0.3; нарративы — до 0.5.',
        'Разработка': 'Температура 0.3–0.5: тесты/инструкции — 0.2–0.3; кейсы/CCAF — 0.4–0.5.',
        'Мета': 'Температура ~0.3: надстройки лучше делать стабильными.'
    }
    for cat, note in recs.items():
        p = doc.add_paragraph()
        r = p.add_run(cat + ': ')
        r.bold = True
        p.add_run(note)

    doc.add_paragraph('')

    # Examples for temperature usage
    doc.add_heading('Примеры настройки температуры', level=1)
    # Example 1
    doc.add_heading('Пример 1: Проверочный тест «Охрана труда»', level=2)
    ex1 = [
        ('0.2', '5 точных вопросов с одним верным вариантом; термины из инструкции; без двусмысленностей.'),
        ('0.4', '7 вопросов, добавь 2 ситуационных; краткая обратная связь после каждого ответа.'),
        ('0.6', '10 вопросов как мини‑сценки с выбором действий; допускай неожиданные, но реалистичные детали.')
    ]
    for temp, text in ex1:
        p = doc.add_paragraph()
        r = p.add_run(f"Температура {temp}: ")
        r.bold = True
        p.add_run(text)
    doc.add_paragraph('')

    # Example 2
    doc.add_heading('Пример 2: Кейс «Работа с жалобой клиента»', level=2)
    ex2 = [
        ('0.2', 'Кейс 150–200 слов, строго по фактам; одна правильная развязка по регламенту.'),
        ('0.4', 'Кейс 400–600 слов, 2 допустимых пути решения; укажи риски каждого.'),
        ('0.6', 'Кейс 800+ слов, конфликт интересов и дефицит времени; 3 правдоподобных разворота, финал открытый.')
    ]
    for temp, text in ex2:
        p = doc.add_paragraph()
        r = p.add_run(f"Температура {temp}: ")
        r.bold = True
        p.add_run(text)
    doc.add_paragraph('')

    cat_order = ['Диагностика', 'Проектирование', 'Коммуникации', 'Разработка', 'Мета']
    for cat in cat_order:
        items = [p for p in protocols if p.get('cat') == cat]
        if not items:
            continue
        doc.add_heading(cat, 1)
        for p in items:
            doc.add_heading(p.get('title', ''), 2)
            d = doc.add_paragraph()
            r = d.add_run('Описание: ')
            r.bold = True
            d.add_run(p.get('desc', ''))
            lab = doc.add_paragraph()
            r = lab.add_run('Промпт:')
            r.bold = True
            for line in p.get('prompt', '').split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
            doc.add_paragraph('')

    doc.save(out_path)


if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    src = os.path.join(base, 'prompts.html')
    out = os.path.join(base, 'docs', 'prompts', 'Prompt_Library_v5.docx')
    # Curate: keep only protocols with уникальная ценность для методистов
    keep_ids = {
        'dp-expert',       # Интервью с экспертом (7 шагов) — уникальный конвейер
        'dp-objections',   # Поведенческая наука в отработке возражений
        'dp-ccaf',         # CCAF+скептик+кейсы — связка разработки
        'dp-ca',           # Ролевое моделирование ЦА с JTBD/маркерами
        'dp-competency',   # Разбор компетенций до наблюдаемых индикаторов
        'dp-meta'          # Мета-техники (панель, слепые зоны, герм. круг)
    }
    prots_all = load_deep_protocols(src)
    prots = [p for p in prots_all if p.get('id') in keep_ids]
    build_docx(prots, out)
    print(f'Saved: {out} ({len(prots)} protocols)')
