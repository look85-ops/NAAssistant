# Собирает .md → .docx в фирменном стиле резюме (синий 1F3A5F, полоса под секциями).
# Запуск: C:\Users\marcenuk\AppData\Local\Python\pythoncore-3.14-64\python.exe scripts/build_resume_docx.py <input.md> <output.docx>

import sys, re
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)

def border_bottom(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3A5F')
    pBdr.append(bottom)
    pPr.append(pBdr)

def run(paragraph, text, bold=False, italic=False, size=10, color=None):
    r = paragraph.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r

def add_runs_split(paragraph, text, size=10, bold_lead=False, bold_all=False):
    """Разбирает **жирный** фрагменты + опциональный жирный lead до первого ':'."""
    if bold_all:
        run(paragraph, text.replace('**', ''), size=size)
        return
    if '**' in text:
        for i, seg in enumerate(text.split('**')):
            if not seg:
                continue
            run(paragraph, seg, bold=(i % 2 == 1), size=size)
        return
    if bold_lead:
        m = re.match(r'^([^:：]+)[:：]', text)
        if m:
            run(paragraph, text[:m.end()], bold=True, size=size)
            run(paragraph, rest_after(text, m.end()), size=size)
            return
    run(paragraph, text, size=size)

def rest_after(text, pos):
    return text[pos:]

def convert(md_path, docx_path):
    doc = Document()

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    first_nonempty = True
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            if first_nonempty:
                continue
            continue

        if line.startswith('# '):
            continue

        if first_nonempty:
            first_nonempty = False
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            run(p, line, bold=True, size=18, color=ACCENT)
            continue

        if 'Телеграм' in line or 'Telegram' in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            run(p, line, size=10)
            continue

        if re.match(r'^(Москва|Moscow)', line.strip()):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            run(p, line, italic=True, size=10)
            continue

        if line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            border_bottom(p)
            run(p, line[3:].strip(), bold=True, size=12, color=ACCENT)
            continue

        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            add_runs_split(p, line[2:], bold_lead=True)
            continue

        if line.startswith('**'):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            add_runs_split(p, line, bold_lead=True)
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_runs_split(p, line, bold_lead=True)

    doc.save(docx_path)
    print(f"OK: {docx_path}")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python build_resume_docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])